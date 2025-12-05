import streamlit as st
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO
from supabase import create_client, Client

# --- Supabase Initialization ---
def init_supabase_client():
    """Initializes and returns the Supabase client if credentials are set."""
    if "supabase" in st.secrets and st.secrets["supabase"]["url"] and st.secrets["supabase"]["key"]:
        supabase_url = st.secrets["supabase"]["url"]
        supabase_key = st.secrets["supabase"]["key"]
        return create_client(supabase_url, supabase_key)
    else:
        st.error("Supabase credentials are not set. Please add them to your Streamlit secrets.")
        return None

supabase: Client = init_supabase_client()

# --- Predefined Locations ---
LOCATIONS = [
    "Chalmers", "Bay View Dental", "Portsoy Medical",
    "Fraserburgh Hospital", "Ugie Hospital", "Peterhead Hospital", "Macduff Vaccination"
]

# --- Session State for Form Inputs ---
if "start_time" not in st.session_state:
    st.session_state.start_time = None
if "air_temp" not in st.session_state:
    st.session_state.air_temp = 0
if "snow" not in st.session_state:
    st.session_state.snow = "No"
if "plough" not in st.session_state:
    st.session_state.plough = "No"
if "grit" not in st.session_state:
    st.session_state.grit = "No"
if "entry_added" not in st.session_state:
    st.session_state.entry_added = False

# Reset form if an entry was just added
if st.session_state.entry_added:
    st.session_state.start_time = None
    st.session_state.air_temp = 0
    st.session_state.snow = "No"
    st.session_state.plough = "No"
    st.session_state.grit = "No"
    st.session_state.entry_added = False
    st.rerun()

st.title("Winter Maintenance Record")

# --- Data Entry Form ---
location = st.selectbox("Select Location", LOCATIONS)

if st.button("Start"):
    st.session_state.start_time = datetime.now().strftime("%H:%M")
    st.rerun()

st.write(f"Entering data for: **{location}**")
start_time_display = st.session_state.start_time if st.session_state.start_time else "Not started"
st.write(f"Entry started at: **{start_time_display}**")

snow_options = ["Yes", "No"]
plough_options = ["Yes", "No"]
grit_options = ["Yes", "No"]

st.number_input("Air Temperature (°C)", key="air_temp", step=1)
st.radio("Snow Present?", snow_options, key="snow", index=snow_options.index(st.session_state.snow))
st.radio("Plough Used?", plough_options, key="plough", index=plough_options.index(st.session_state.plough))
st.radio("Grit Spread?", grit_options, key="grit", index=grit_options.index(st.session_state.grit))

if st.button("Add Entry"):
    if st.session_state.start_time is None:
        st.error("Please click 'Start' before adding an entry.")
    elif supabase:
        finish_time = datetime.now().strftime("%H:%M")
        date_today = datetime.now().strftime("%d-%m-%Y")

        new_entry = {
            "date": date_today,
            "start_time": st.session_state.start_time,
            "finish_time": finish_time,
            "location": location,
            "air_temp": st.session_state.air_temp,
            "snow": st.session_state.snow,
            "plough": st.session_state.plough,
            "grit": st.session_state.grit
        }
        
        try:
            supabase.table("winter_maintenance").insert(new_entry).execute()
            st.session_state.entry_added = True
            st.success("Entry added successfully to Supabase!")
        except Exception as e:
            st.error(f"Failed to add entry to Supabase: {e}")

# --- Display Entries and Download Option ---
if supabase:
    try:
        response = supabase.table("winter_maintenance").select("*").order("created_at", desc=True).execute()
        entries = response.data
        
        if entries:
            st.subheader("Current Entries from Supabase")
            # Prepare data for display (e.g., format dates, select columns)
            display_entries = [
                {
                    "Date": e["date"], "Start Time": e["start_time"], "Finish Time": e["finish_time"],
                    "Location": e["location"], "Air Temp": e["air_temp"], "Snow": e["snow"],
                    "Plough": e["plough"], "Grit": e["grit"]
                } for e in entries
            ]
            st.table(display_entries)

            # --- Word Document Creation ---
            def create_word(entries_data):
                doc = Document()
                doc.add_heading("Winter Maintenance Record", 0)
                table = doc.add_table(rows=1, cols=8)
                table.columns[3].width = Inches(1.2)
                hdr_cells = table.rows[0].cells
                headers = ['Date', 'Start Time', 'Finish Time', 'Location', 'Air Temp', 'Snow', 'Plough Used', 'Grit Spread']
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                
                for entry in entries_data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = entry["date"]
                    row_cells[1].text = entry["start_time"]
                    row_cells[2].text = entry["finish_time"]
                    p = row_cells[3].paragraphs[0]
                    run = p.add_run(entry["location"])
                    run.font.size = Pt(8)
                    row_cells[4].text = str(entry["air_temp"])
                    row_cells[5].text = entry["snow"]
                    row_cells[6].text = entry["plough"]
                    row_cells[7].text = entry["grit"]

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer

            word_file = create_word(entries)
            
            # --- Download and Clear Buttons ---
            col1, col2 = st.columns(2)
            
            with col1:
                st.download_button(
                    label="Download as Word",
                    data=word_file,
                    file_name=f"grit_report_{datetime.now().strftime('%Y-%m-%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            with col2:
                if st.button("Clear All Entries from Database"):
                    try:
                        # In a real-world app, you might want a confirmation modal here
                        # For simplicity, we directly delete.
                        supabase.table("winter_maintenance").delete().gt("id", 0).execute()
                        st.success("All entries have been cleared from the database!")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Failed to clear entries: {e}")

    except Exception as e:
        st.error(f"Error fetching data from Supabase: {e}")
        st.info("Please ensure your Supabase table 'winter_maintenance' is created and credentials are correct.")